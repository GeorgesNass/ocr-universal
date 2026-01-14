'''
__author__ = "Georges Nassopoulos"
__copyright__ = None
__version__ = "1.0.0"
__email__ = "georges.nassopoulos@gmail.com"
__status__ = "Prod"
__desc__ = "Convert DOC to DOCX and extract text from DOCX files using LibreOffice and python-docx, with optional headers and tables."
'''

import os
import io
import sys
import time
import docx
from pathlib import Path
from typing import Optional
import subprocess

from src.utils import get_logger, log_execution_time_and_path
from src.utils.constants import (
    BASE_DIR,
    INPUT_DIR,
    CONVERTED_DIR,
    OUTPUT_DIR,
    ALLOWED_EXTENSIONS,
    DIR_SEPARATOR,
    PATH_LIBRE_OFFICE,
    INCLUDE_DOCX_HEADERS,
    INCLUDE_DOCX_TABLES
)

## ============================================================
## LOGGER INITIALIZATION
## ============================================================

logger = get_logger("docx_doc_to_text")

## ============================================================
## DOC TO DOCX CONVERSION
## ============================================================
@log_execution_time_and_path
def convert_doc_to_docx(src_path: str) -> Optional[str]:
    """
        Convert a .doc file to .docx format using LibreOffice in headless mode

        Args:
            src_path (str): File path complet to the file

        Returns:
            Optional[str]: Path to the converted .docx file, or None if conversion failed
    """
    
    ## Keep homogeneous exists check (like PDF)
    if not os.path.exists(str(src_path)):
        logger.error(f"Source file not found: {src_path}")
        return None

    ## Normalize to Path ONLY for suffix/stem building (do not rebuild with INPUT_DIR)
    src_path = Path(src_path)

    ## Build expected output file path
    docx_path = CONVERTED_DIR / f"{src_path.stem}.docx"

    ## Build LibreOffice conversion command
    cmd = [
        str(PATH_LIBRE_OFFICE),
        "--headless",
        "--convert-to", "docx",
        str(src_path),
        "--outdir", str(CONVERTED_DIR)
    ]
    logger.info(f"Executing LibreOffice command: {' '.join(cmd)}")

    try:
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            logger.error(f"LibreOffice conversion failed: {result.stderr}")
            return None
    except Exception as e:
        logger.exception(f"Error converting DOC to DOCX: {e}")
        return None

    if not os.path.exists(str(docx_path)):
        logger.error(f"Conversion failed: {docx_path} not found.")
        return None

    logger.info(f"File converted successfully: {docx_path}")
    return str(docx_path)

## ============================================================
## DOCX TEXT EXTRACTION - SIMPLE VERSION
## ============================================================
def extract_text_simple(file_path: str) -> str:
    """
        Extract plain text from DOCX (paragraphs only)

        Args:
            file_path (str): Path to the DOCX file

        Returns:
            str: Extracted text without headers or tables
    """
    
    fulltext = []
    try:
        document = docx.Document(file_path)
        logger.debug(f"Opened document: {file_path}")

        ## Extract paragraph text
        for para in document.paragraphs:
            fulltext.append(para.text)

    except Exception as e:
        logger.exception(f"Error reading {file_path}: {e}")

    return "\n".join(fulltext)

## ============================================================
## DOCX TEXT EXTRACTION - DETAILED VERSION
## ============================================================
def extract_text_detailed(file_path: str) -> str:
    """
        Extract DOCX content including headers and tables if enabled

        Args:
            file_path (str): Path to the DOCX file

        Returns:
            str: Extracted text with optional headers and tables
    """
    
    fulltext = []
    
    try:
        document = docx.Document(file_path)
        logger.debug(f"Opened document: {file_path}")

        ## Include header text if enabled
        if INCLUDE_DOCX_HEADERS and hasattr(document, "sections"):
            for section in document.sections:
                header = section.header
                for para in header.paragraphs:
                    fulltext.append(para.text)
            logger.debug("Included DOCX headers")

        ## Include table text if enabled
        if INCLUDE_DOCX_TABLES:
            for table in document.tables:
                for row in table.rows:
                    row_text = "\t".join(
                        " ".join(p.text for p in cell.paragraphs)
                        for cell in row.cells
                    )
                    fulltext.append(row_text)
                fulltext.append("\n")
            logger.debug("Included DOCX tables")

        ## Extract main body paragraphs
        for para in document.paragraphs:
            fulltext.append(para.text)

    except Exception as e:
        logger.exception(f"Error reading {file_path}: {e}")

    return "\n".join(fulltext)

## ============================================================
## MASTER FUNCTION: SELECT EXTRACTION MODE
## ============================================================
@log_execution_time_and_path
def get_text_from_docx(file_path: str) -> str:
    """
        Main entry to extract text from DOCX depending on global flags

        Args:
            file_path (str): Path to the DOCX file

        Returns:
            str: Extracted text
    """
    
    try:
        if INCLUDE_DOCX_HEADERS or INCLUDE_DOCX_TABLES:
            logger.info("Using detailed DOCX extraction (headers/tables enabled)")
            return extract_text_detailed(file_path)
        else:
            logger.info("Using simple DOCX extraction (paragraphs only)")
            return extract_text_simple(file_path)
    except Exception as e:
        logger.exception(f"Failed to extract text: {e}")
        return ""

## ============================================================
## SAVE OUTPUT
## ============================================================
@log_execution_time_and_path
def save_extracted_text(file_path: Path, text: str) -> str:
    """
        Save extracted text to the output directory

        Args:
            file_path (Path): Original DOCX file path
            text (str): Extracted text

        Returns:
            str: Path to saved output file
    """
    
    output_path = OUTPUT_DIR / f"{file_path.stem}.txt"
    try:
        output_path.write_text(text, encoding="utf-8", errors="ignore")
        logger.info(f"Saved extracted text to: {output_path}")
        return str(output_path)
    except Exception as e:
        logger.exception(f"Error saving text to {output_path}: {e}")
        return ""

## ============================================================
## PIPELINE: FULL EXTRACTION PROCESS
## ============================================================
@log_execution_time_and_path
def process_doc_or_docx(src_path: str) -> None:
    """
        Full pipeline to process DOC or DOCX files:
            - Convert .doc to .docx if needed
            - Extract text (simple or detailed)
            - Save result in /data/output

        Args:
            src_path (str): File path complet to process
    """
    
    ## Keep homogeneous exists check (like PDF)
    if not os.path.exists(str(src_path)):
        logger.error(f"File not found: {src_path}")
        return

    ## Normalize to Path only for suffix handling
    src_path = Path(src_path)

    ## Convert DOC → DOCX if needed
    if src_path.suffix.lower() == ".doc":
        logger.info(f"Converting DOC to DOCX: {src_path}")
        converted = convert_doc_to_docx(str(src_path))
        if not converted:
            return
        src_path = Path(converted)

    ## Extract text
    text = get_text_from_docx(str(src_path))

    ## Save to output
    save_extracted_text(src_path, text)

## ============================================================
## MAIN ENTRY POINT
## ============================================================
if __name__ == "__main__":
    """
        Example manual execution for DOC/DOCX extraction
    """
    
    logger.info("Starting manual execution: DOC/DOCX text extraction")

    ## Defensive: ensure INPUT_DIR exists
    if not INPUT_DIR.exists():
        logger.error(f"INPUT_DIR does not exist: {INPUT_DIR}")
        raise SystemExit(1)

    ## Process only DOC/DOCX
    for file in INPUT_DIR.glob("*"):
        if not file.is_file():
            continue

        if file.suffix.lower() in [".doc", ".docx"]:
            logger.info(f"Processing file: {file}")
            try:
                ## Ensure str is passed (process_doc_or_docx expects str)
                process_doc_or_docx(str(file))
            except Exception as e:
                logger.exception(f"Failed processing file {file}: {e}")

    logger.info("Finished manual processing of DOC/DOCX files.")